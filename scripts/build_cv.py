from datetime import datetime
from pathlib import Path
from zoneinfo import ZoneInfo
import re

import yaml
from docx import Document
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ==========================================================
# Paths
# ==========================================================
ROOT = Path(__file__).resolve().parents[1]

CV_DATA = ROOT / "_data" / "cv.yml"
PUBLICATIONS_DIR = ROOT / "_publications"
LOGO = ROOT / "images" / "unist-logo.png"

OUTPUT_DIR = ROOT / "files"
OUTPUT_DOCX = OUTPUT_DIR / "Nakbin_Choi_CV.docx"


# ==========================================================
# Personal information
# ==========================================================
NAME = "Nakbin Choi, Ph.D."
POSITION = "Research Assistant Professor"

DEPARTMENT = (
    "Department of Civil, Urban, Earth, "
    "and Environmental Engineering"
)

INSTITUTION = (
    "Ulsan National Institute of Science "
    "and Technology (UNIST)"
)

ADDRESS = (
    "50 UNIST-gil, Ulju-gun, Ulsan 44919, "
    "Republic of Korea"
)

EMAIL = "nbchoi21@unist.ac.kr"

PHONE = (
    "(Korea) 010-4946-6138  |  "
    "(US) +1 (703) 843-4060"
)


OUTPUT_DIR.mkdir(
    parents=True,
    exist_ok=True,
)


# ==========================================================
# Font
# ==========================================================
def set_run(
    run,
    size=10.0,
    bold=False,
    italic=False,
    color="000000",
):

    run.font.name = "Times New Roman"

    run._element.get_or_add_rPr().rFonts.set(
        qn("w:eastAsia"),
        "Times New Roman",
    )

    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


# ==========================================================
# Table formatting
# ==========================================================
def set_cell_margins(
    cell,
    top=40,
    start=0,
    bottom=40,
    end=0,
):

    cell_properties = cell._tc.get_or_add_tcPr()

    cell_margins = (
        cell_properties.first_child_found_in(
            "w:tcMar"
        )
    )

    if cell_margins is None:

        cell_margins = OxmlElement("w:tcMar")
        cell_properties.append(cell_margins)

    margin_values = [
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ]

    for name, value in margin_values:

        margin = cell_margins.find(
            qn(f"w:{name}")
        )

        if margin is None:

            margin = OxmlElement(
                f"w:{name}"
            )

            cell_margins.append(margin)

        margin.set(
            qn("w:w"),
            str(value),
        )

        margin.set(
            qn("w:type"),
            "dxa",
        )


def remove_table_borders(table):

    table_properties = table._tbl.tblPr

    borders = (
        table_properties.first_child_found_in(
            "w:tblBorders"
        )
    )

    if borders is None:

        borders = OxmlElement(
            "w:tblBorders"
        )

        table_properties.append(borders)

    edges = [
        "top",
        "left",
        "bottom",
        "right",
        "insideH",
        "insideV",
    ]

    for edge in edges:

        border = borders.find(
            qn(f"w:{edge}")
        )

        if border is None:

            border = OxmlElement(
                f"w:{edge}"
            )

            borders.append(border)

        border.set(
            qn("w:val"),
            "nil",
        )


def set_table_indent(
    table,
    indent=120,
):

    table_properties = table._tbl.tblPr

    table_indent = (
        table_properties.first_child_found_in(
            "w:tblInd"
        )
    )

    if table_indent is None:

        table_indent = OxmlElement(
            "w:tblInd"
        )

        table_properties.append(
            table_indent
        )

    table_indent.set(
        qn("w:w"),
        str(indent),
    )

    table_indent.set(
        qn("w:type"),
        "dxa",
    )


# ==========================================================
# Paragraph formatting
# ==========================================================
def add_bottom_border(
    paragraph,
    size=6,
):

    paragraph_properties = (
        paragraph._p.get_or_add_pPr()
    )

    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")

    bottom.set(
        qn("w:val"),
        "single",
    )

    bottom.set(
        qn("w:sz"),
        str(size),
    )

    bottom.set(
        qn("w:space"),
        "1",
    )

    bottom.set(
        qn("w:color"),
        "000000",
    )

    borders.append(bottom)

    paragraph_properties.append(
        borders
    )


def add_section_heading(
    document,
    title,
):

    paragraph = document.add_paragraph()

    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True

    set_run(
        paragraph.add_run(
            title.upper()
        ),
        size=11.5,
        bold=True,
    )

    add_bottom_border(paragraph)


def add_bullet(
    document,
    text,
):

    paragraph = document.add_paragraph()

    paragraph.paragraph_format.left_indent = (
        Inches(0.22)
    )

    paragraph.paragraph_format.first_line_indent = (
        Inches(-0.16)
    )

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)

    set_run(
        paragraph.add_run(
            "•  " + text
        ),
        size=9.6,
    )


# ==========================================================
# Timeline sections
# ==========================================================
def add_timeline(
    document,
    entries,
):

    table = document.add_table(
        rows=0,
        cols=2,
    )

    table.alignment = (
        WD_TABLE_ALIGNMENT.LEFT
    )

    table.autofit = False

    table.columns[0].width = Inches(1.55)
    table.columns[1].width = Inches(5.43)

    remove_table_borders(table)

    set_table_indent(
        table,
        indent=120,
    )

    for entry in entries:

        cells = table.add_row().cells

        cells[0].width = Inches(1.55)
        cells[1].width = Inches(5.43)

        for cell in cells:

            cell.vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.TOP
            )

            set_cell_margins(
                cell,
                top=35,
                start=0,
                bottom=110,
                end=0,
            )

        date_paragraph = (
            cells[0].paragraphs[0]
        )

        date_paragraph.paragraph_format.left_indent = (
            Inches(0)
        )

        date_paragraph.paragraph_format.first_line_indent = (
            Inches(0)
        )

        date_paragraph.paragraph_format.space_before = Pt(0)
        date_paragraph.paragraph_format.space_after = Pt(0)

        set_run(
            date_paragraph.add_run(
                str(
                    entry.get(
                        "year",
                        "",
                    )
                )
            ),
            size=9.4,
        )

        content_paragraph = (
            cells[1].paragraphs[0]
        )

        content_paragraph.paragraph_format.left_indent = (
            Inches(0)
        )

        content_paragraph.paragraph_format.first_line_indent = (
            Inches(0)
        )

        content_paragraph.paragraph_format.space_before = Pt(0)
        content_paragraph.paragraph_format.space_after = Pt(0)
        content_paragraph.paragraph_format.keep_together = True

        set_run(
            content_paragraph.add_run(
                str(
                    entry.get(
                        "title",
                        "",
                    )
                )
            ),
            size=9.6,
            bold=True,
        )

        institution = str(
            entry.get(
                "institution",
                "",
            )
        ).strip()

        if institution:

            set_run(
                content_paragraph.add_run(
                    ", " + institution
                ),
                size=9.6,
            )

        description = entry.get(
            "description",
            "",
        )

        if isinstance(
            description,
            list,
        ):

            description_lines = [
                str(item)
                for item in description
            ]

        elif description:

            description_lines = [
                str(description)
            ]

        else:

            description_lines = []

        for line in description_lines:

            paragraph = (
                cells[1].add_paragraph()
            )

            paragraph.paragraph_format.left_indent = (
                Inches(0)
            )

            paragraph.paragraph_format.first_line_indent = (
                Inches(0)
            )

            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.keep_together = True

            if ":" in line:

                label, value = line.split(
                    ":",
                    1,
                )

                set_run(
                    paragraph.add_run(
                        label + ": "
                    ),
                    size=9.2,
                    bold=True,
                )

                set_run(
                    paragraph.add_run(
                        value.strip()
                    ),
                    size=9.2,
                    italic=(
                        label.lower()
                        == "dissertation"
                    ),
                )

            else:

                set_run(
                    paragraph.add_run(line),
                    size=9.2,
                )


# ==========================================================
# Publication data
# ==========================================================
def read_frontmatter(path):

    text = path.read_text(
        encoding="utf-8"
    )

    match = re.match(
        r"^---\s*\n(.*?)\n---",
        text,
        re.DOTALL,
    )

    if match is None:
        return {}

    return (
        yaml.safe_load(
            match.group(1)
        )
        or {}
    )


def add_publication(
    document,
    number,
    citation,
):

    paragraph = document.add_paragraph()

    paragraph.paragraph_format.left_indent = (
        Inches(0.31)
    )

    paragraph.paragraph_format.first_line_indent = (
        Inches(-0.31)
    )

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_together = True

    set_run(
        paragraph.add_run(
            f"[{number}] "
        ),
        size=9.5,
    )

    bold_tokens = [
        "Choi, N.",
        "Nakbin Choi",
    ]

    position = 0

    while position < len(citation):

        matches = [
            (
                citation.find(
                    token,
                    position,
                ),
                token,
            )
            for token in bold_tokens
        ]

        matches = [
            (index, token)
            for index, token in matches
            if index >= 0
        ]

        if not matches:

            set_run(
                paragraph.add_run(
                    citation[position:]
                ),
                size=9.5,
            )

            break

        index, token = min(matches)

        if index > position:

            set_run(
                paragraph.add_run(
                    citation[
                        position:index
                    ]
                ),
                size=9.5,
            )

        set_run(
            paragraph.add_run(token),
            size=9.5,
            bold=True,
        )

        position = (
            index
            + len(token)
        )


# ==========================================================
# Read CV data
# ==========================================================
with CV_DATA.open(
    "r",
    encoding="utf-8",
) as file:

    cv_sections = (
        yaml.safe_load(file)
        or []
    )


publications = []

for path in PUBLICATIONS_DIR.glob(
    "*.md"
):

    metadata = read_frontmatter(path)

    citation = str(
        metadata.get(
            "citation",
            "",
        )
    ).strip()

    if not citation:
        continue

    if not citation.endswith("."):
        citation += "."

    publications.append(
        {
            "date": str(
                metadata.get(
                    "date",
                    "",
                )
            ),
            "citation": citation,
        }
    )


publications.sort(
    key=lambda item: item["date"],
    reverse=True,
)


# ==========================================================
# Create document
# ==========================================================
document = Document()

section = document.sections[0]

section.page_width = Inches(8.5)
section.page_height = Inches(11)

section.top_margin = Inches(0.58)
section.bottom_margin = Inches(0.58)
section.left_margin = Inches(0.72)
section.right_margin = Inches(0.72)

section.header_distance = Inches(0.28)
section.footer_distance = Inches(0.28)


normal_style = document.styles[
    "Normal"
]

normal_style.font.name = (
    "Times New Roman"
)

normal_style._element.rPr.rFonts.set(
    qn("w:eastAsia"),
    "Times New Roman",
)

normal_style.font.size = Pt(10)

normal_style.paragraph_format.space_before = Pt(0)
normal_style.paragraph_format.space_after = Pt(0)
normal_style.paragraph_format.line_spacing = 1.0


# ==========================================================
# Footer
# ==========================================================
footer = section.footer

footer_paragraph = (
    footer.paragraphs[0]
)

footer_paragraph.paragraph_format.tab_stops.add_tab_stop(
    Inches(6.65)
)

set_run(
    footer_paragraph.add_run(NAME),
    size=8.5,
)

set_run(
    footer_paragraph.add_run("\t"),
    size=8.5,
)

page_field = OxmlElement(
    "w:fldSimple"
)

page_field.set(
    qn("w:instr"),
    "PAGE",
)

footer_paragraph._p.append(
    page_field
)

add_bottom_border(
    footer_paragraph
)


# ==========================================================
# Title
# ==========================================================
paragraph = document.add_paragraph()

paragraph.alignment = (
    WD_ALIGN_PARAGRAPH.CENTER
)

paragraph.paragraph_format.space_after = Pt(0)

set_run(
    paragraph.add_run(
        "CURRICULUM VITAE"
    ),
    size=24,
    bold=True,
)


paragraph = document.add_paragraph()

paragraph.alignment = (
    WD_ALIGN_PARAGRAPH.RIGHT
)

paragraph.paragraph_format.space_after = Pt(2)

cv_date = datetime.now(
    ZoneInfo("Asia/Seoul")
).strftime("%B %Y")

set_run(
    paragraph.add_run(cv_date),
    size=9.0,
    bold=True,
)

add_bottom_border(
    paragraph,
    size=10,
)


# ==========================================================
# Contact block
# ==========================================================
contact_table = document.add_table(
    rows=1,
    cols=2,
)

contact_table.alignment = (
    WD_TABLE_ALIGNMENT.LEFT
)

contact_table.autofit = False

contact_table.columns[0].width = (
    Inches(5.80)
)

contact_table.columns[1].width = (
    Inches(1.18)
)

remove_table_borders(
    contact_table
)

set_table_indent(
    contact_table,
    indent=120,
)

left_cell, right_cell = (
    contact_table.rows[0].cells
)

left_cell.width = Inches(5.80)
right_cell.width = Inches(1.18)

set_cell_margins(
    left_cell,
    top=70,
    start=0,
    bottom=75,
    end=0,
)

set_cell_margins(
    right_cell,
    top=20,
    start=0,
    bottom=40,
    end=0,
)


paragraph = (
    left_cell.paragraphs[0]
)

paragraph.paragraph_format.left_indent = (
    Inches(0)
)

paragraph.paragraph_format.first_line_indent = (
    Inches(0)
)

paragraph.paragraph_format.space_after = Pt(8)

set_run(
    paragraph.add_run(NAME),
    size=17,
    bold=True,
)


contact_lines = [
    (
        POSITION,
        True,
        False,
    ),
    (
        DEPARTMENT,
        False,
        True,
    ),
    (
        INSTITUTION,
        False,
        True,
    ),
    (
        ADDRESS,
        False,
        False,
    ),
]

for text, bold, italic in contact_lines:

    paragraph = (
        left_cell.add_paragraph()
    )

    paragraph.paragraph_format.left_indent = (
        Inches(0)
    )

    paragraph.paragraph_format.first_line_indent = (
        Inches(0)
    )

    paragraph.paragraph_format.space_after = Pt(
        1.5
    )

    set_run(
        paragraph.add_run(text),
        size=9.4,
        bold=bold,
        italic=italic,
    )


paragraph = left_cell.add_paragraph()

paragraph.paragraph_format.left_indent = (
    Inches(0)
)

paragraph.paragraph_format.first_line_indent = (
    Inches(0)
)

paragraph.paragraph_format.space_after = Pt(
    1.5
)

email_run = paragraph.add_run(EMAIL)

set_run(
    email_run,
    size=9.4,
    color="0000CC",
)

email_run.underline = True


paragraph = left_cell.add_paragraph()

paragraph.paragraph_format.left_indent = (
    Inches(0)
)

paragraph.paragraph_format.first_line_indent = (
    Inches(0)
)

set_run(
    paragraph.add_run(PHONE),
    size=9.4,
)


logo_paragraph = (
    right_cell.paragraphs[0]
)

logo_paragraph.alignment = (
    WD_ALIGN_PARAGRAPH.RIGHT
)

logo_paragraph.paragraph_format.left_indent = (
    Inches(0)
)

logo_paragraph.paragraph_format.right_indent = (
    Inches(0)
)

logo_paragraph.paragraph_format.first_line_indent = (
    Inches(0)
)

logo_paragraph.paragraph_format.space_before = Pt(0)
logo_paragraph.paragraph_format.space_after = Pt(0)

logo_paragraph.add_run().add_picture(
    str(LOGO),
    width=Inches(0.92),
)


# ==========================================================
# CV sections from _data/cv.yml
# ==========================================================
for cv_section in cv_sections:

    title = str(
        cv_section.get(
            "title",
            "",
        )
    ).strip()

    section_type = str(
        cv_section.get(
            "type",
            "list",
        )
    ).strip()

    contents = (
        cv_section.get(
            "contents",
            [],
        )
        or []
    )

    if not title or not contents:
        continue

    add_section_heading(
        document,
        title,
    )

    if section_type == "time_table":

        add_timeline(
            document,
            contents,
        )

    else:

        for item in contents:

            add_bullet(
                document,
                str(item),
            )


# ==========================================================
# Publications
# ==========================================================
if publications:

    add_section_heading(
        document,
        "Publications",
    )

    publication_count = len(
        publications
    )

    for index, publication in enumerate(
        publications
    ):

        add_publication(
            document,
            publication_count - index,
            publication["citation"],
        )


# ==========================================================
# Metadata and output
# ==========================================================
document.core_properties.title = (
    "Curriculum Vitae - Nakbin Choi"
)

document.core_properties.subject = (
    "Academic curriculum vitae"
)

document.core_properties.author = (
    "Nakbin Choi"
)

document.save(OUTPUT_DOCX)

print(f"Created: {OUTPUT_DOCX}")
